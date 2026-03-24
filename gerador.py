from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re

def formatar_paragrafo(doc, texto):
    """
    Detecta se o texto é um título ou parágrafo normal
    e formata corretamente no Word, removendo Markdown.
    """

    # Remove asteriscos do Markdown (**texto** ou ***texto***)
    texto_limpo = re.sub(r'\*+', '', texto).strip()

    if not texto_limpo:
        return

    # Detecta títulos numerados: "1.", "2.", "1.1", etc
    eh_titulo = bool(re.match(r'^\d+[\.\d]*\s', texto_limpo))

    # Detecta títulos com "#"
    if texto_limpo.startswith("#"):
        texto_limpo = texto_limpo.lstrip("#").strip()
        eh_titulo = True

    if eh_titulo:
        doc.add_heading(texto_limpo, level=2)
    else:
        doc.add_paragraph(texto_limpo)


def gerar_relatorio(analise, dados, tipo_relatorio, caminho_saida):
    """
    Recebe o texto da IA e os dados da planilha e gera um arquivo .docx
    """

    doc = Document()

    # ── Título principal ───────────────────────────
    titulo = doc.add_heading(f"Relatório de {tipo_relatorio}", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Data de geração ────────────────────────────
    data_atual = datetime.now().strftime("%d/%m/%Y às %H:%M")
    subtitulo = doc.add_paragraph(f"Gerado automaticamente em {data_atual}")
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ── Informações gerais da planilha ─────────────
    doc.add_heading("Informações da Planilha", level=1)
    doc.add_paragraph(f"Total de registros: {dados['total_linhas']}")
    doc.add_paragraph(f"Total de colunas: {dados['total_colunas']}")
    doc.add_paragraph(f"Colunas: {', '.join(dados['colunas'])}")

    doc.add_paragraph()

    # ── Resumo estatístico em tabela ───────────────
    doc.add_heading("Resumo Estatístico", level=1)

    tabela = doc.add_table(rows=1, cols=5)
    tabela.style = "Table Grid"

    cabecalho = tabela.rows[0].cells
    cabecalho[0].text = "Coluna"
    cabecalho[1].text = "Média"
    cabecalho[2].text = "Mínimo"
    cabecalho[3].text = "Máximo"
    cabecalho[4].text = "Total"

    for coluna, stats in dados["resumo_estatistico"].items():
        linha = tabela.add_row().cells
        linha[0].text = coluna
        linha[1].text = str(stats["média"])
        linha[2].text = str(stats["mínimo"])
        linha[3].text = str(stats["máximo"])
        linha[4].text = str(stats["total"])

    doc.add_paragraph()

    # ── Análise da IA ──────────────────────────────
    doc.add_heading("Análise Gerada por IA", level=1)

    for linha in analise.split("\n"):
        formatar_paragrafo(doc, linha)

    # ── Salva o arquivo ────────────────────────────
    doc.save(caminho_saida)
    return caminho_saida