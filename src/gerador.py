from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re


def formatar_paragrafo(doc, texto):
    texto_limpo = re.sub(r'\*+', '', texto).strip()

    if not texto_limpo:
        return

    if texto_limpo.startswith("#"):
        texto_limpo = texto_limpo.lstrip("#").strip()
        doc.add_heading(texto_limpo, level=2)
        return

    eh_titulo_simples = bool(re.match(r'^\d+\.\s+\S', texto_limpo)) and len(texto_limpo) < 60

    if eh_titulo_simples:
        doc.add_heading(texto_limpo, level=2)
    else:
        doc.add_paragraph(texto_limpo)


def gerar_relatorio(analise, dados, tipo_relatorio, idioma, caminho_saida):

    doc = Document()
    data_atual = datetime.now().strftime("%d/%m/%Y às %H:%M")

    # ── Traduções ──────────────────────────────────
    traducoes = {
        "Português": {
            "titulo":    f"Relatório de {tipo_relatorio}",
            "gerado":    f"Gerado automaticamente em {data_atual}",
            "info":      "Informações da Planilha",
            "registros": f"Total de registros: {dados['total_linhas']}",
            "colunas_n": f"Total de colunas: {dados['total_colunas']}",
            "colunas":   f"Colunas: {', '.join(dados['colunas'])}",
            "resumo":    "Resumo Estatístico",
            "analise":   "Análise Gerada por IA",
        },
        "Inglês": {
            "titulo":    f"{tipo_relatorio} Report",
            "gerado":    f"Automatically generated on {data_atual}",
            "info":      "Spreadsheet Information",
            "registros": f"Total records: {dados['total_linhas']}",
            "colunas_n": f"Total columns: {dados['total_colunas']}",
            "colunas":   f"Columns: {', '.join(dados['colunas'])}",
            "resumo":    "Statistical Summary",
            "analise":   "AI Generated Analysis",
        },
        "Espanhol": {
            "titulo":    f"Informe de {tipo_relatorio}",
            "gerado":    f"Generado automáticamente el {data_atual}",
            "info":      "Información de la Hoja de Cálculo",
            "registros": f"Total de registros: {dados['total_linhas']}",
            "colunas_n": f"Total de columnas: {dados['total_colunas']}",
            "colunas":   f"Columnas: {', '.join(dados['colunas'])}",
            "resumo":    "Resumen Estadístico",
            "analise":   "Análisis Generado por IA",
        },
    }

    t = traducoes.get(idioma, traducoes["Português"])

    # ── Título principal ───────────────────────────
    titulo = doc.add_heading(t["titulo"], level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Data de geração ────────────────────────────
    subtitulo = doc.add_paragraph(t["gerado"])
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ── Informações gerais da planilha ─────────────
    doc.add_heading(t["info"], level=1)
    doc.add_paragraph(t["registros"])
    doc.add_paragraph(t["colunas_n"])
    doc.add_paragraph(t["colunas"])

    doc.add_paragraph()

    # ── Resumo estatístico em tabela ───────────────
    doc.add_heading(t["resumo"], level=1)

    tabela = doc.add_table(rows=1, cols=5)
    tabela.style = "Table Grid"

    cabecalho = tabela.rows[0].cells
    cabecalho[0].text = "Column"
    cabecalho[1].text = "Mean"
    cabecalho[2].text = "Min"
    cabecalho[3].text = "Max"
    cabecalho[4].text = "Sum"

    for coluna, stats in dados["resumo_estatistico"].items():
        linha = tabela.add_row().cells
        linha[0].text = coluna
        linha[1].text = str(stats["mean"])
        linha[2].text = str(stats["min"])
        linha[3].text = str(stats["max"])
        linha[4].text = str(stats["sum"])

    doc.add_paragraph()

    # ── Análise da IA ──────────────────────────────
    doc.add_heading(t["analise"], level=1)

    for linha in analise.split("\n"):
        formatar_paragrafo(doc, linha)

    # ── Salva o arquivo ────────────────────────────
    doc.save(caminho_saida)
    return caminho_saida